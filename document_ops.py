import os
import subprocess
import platform
from datetime import datetime
import tempfile
import shutil

# Try to import document processing libraries
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

# ==================== DOCUMENT CREATION ====================

def create_word_document(filename=None, content="", location=None):
    """Create a new Word document"""
    try:
        if not DOCX_AVAILABLE:
            return "Error: python-docx library not installed. Run: pip install python-docx"
        
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"document_{timestamp}.docx"
        
        if not filename.endswith('.docx'):
            filename += '.docx'
        
        if location is None:
            location = os.path.join(os.path.expanduser("~"), "Documents")
            if not os.path.exists(location):
                location = os.path.expanduser("~")
        
        full_path = os.path.join(location, filename)
        
        # Create document
        doc = Document()
        
        if content:
            doc.add_paragraph(content)
        else:
            doc.add_heading('New Document', 0)
            doc.add_paragraph('This is a new document created by Desktop AI.')
            doc.add_paragraph('You can edit this document in Microsoft Word or any compatible word processor.')
        
        doc.save(full_path)
        
        return f"Word document created: {full_path}"
    
    except Exception as e:
        return f"Error creating Word document: {str(e)}"

def open_word_document(filepath):
    """Open a Word document in the default application"""
    try:
        if not os.path.exists(filepath):
            return f"File not found: {filepath}"
        
        if platform.system().lower() == "windows":
            os.startfile(filepath)
        elif platform.system().lower() == "darwin":  # macOS
            subprocess.run(["open", filepath])
        else:  # Linux
            subprocess.run(["xdg-open", filepath])
        
        return f"Opened document: {filepath}"
    
    except Exception as e:
        return f"Error opening document: {str(e)}"

# ==================== DOCUMENT CONVERSION ====================

def convert_docx_to_pdf(docx_path, pdf_path=None):
    """Convert DOCX to PDF"""
    try:
        if not os.path.exists(docx_path):
            return f"DOCX file not found: {docx_path}"
        
        if pdf_path is None:
            pdf_path = docx_path.replace('.docx', '.pdf')
        
        if platform.system().lower() == "windows":
            # Try using docx2pdf library first
            if DOCX2PDF_AVAILABLE:
                try:
                    convert(docx_path, pdf_path)
                    return f"Successfully converted to PDF: {pdf_path}"
                except Exception as e:
                    pass
            
            # Fallback to LibreOffice if available
            try:
                subprocess.run([
                    "soffice", "--headless", "--convert-to", "pdf", 
                    "--outdir", os.path.dirname(pdf_path), docx_path
                ], check=True, capture_output=True)
                return f"Successfully converted to PDF: {pdf_path}"
            except (subprocess.CalledProcessError, FileNotFoundError):
                return "PDF conversion failed. Please install LibreOffice or python-docx2pdf library."
        
        else:
            # Linux/Mac - try LibreOffice
            try:
                subprocess.run([
                    "libreoffice", "--headless", "--convert-to", "pdf", 
                    "--outdir", os.path.dirname(pdf_path), docx_path
                ], check=True, capture_output=True)
                return f"Successfully converted to PDF: {pdf_path}"
            except (subprocess.CalledProcessError, FileNotFoundError):
                return "PDF conversion failed. Please install LibreOffice."
    
    except Exception as e:
        return f"Error converting DOCX to PDF: {str(e)}"

def convert_pdf_to_docx(pdf_path, docx_path=None):
    """Convert PDF to DOCX (basic text extraction)"""
    try:
        if not PDF_AVAILABLE:
            return "Error: PDF processing libraries not installed. Run: pip install PyPDF2 pdfplumber"
        
        if not os.path.exists(pdf_path):
            return f"PDF file not found: {pdf_path}"
        
        if docx_path is None:
            docx_path = pdf_path.replace('.pdf', '.docx')
        
        if not DOCX_AVAILABLE:
            return "Error: python-docx library not installed. Run: pip install python-docx"
        
        # Extract text from PDF
        text_content = extract_text_from_pdf(pdf_path)
        if text_content.startswith("Error"):
            return text_content
        
        # Create DOCX with extracted text
        doc = Document()
        doc.add_heading('Converted from PDF', 0)
        doc.add_paragraph(text_content)
        doc.save(docx_path)
        
        return f"Successfully converted to DOCX: {docx_path}"
    
    except Exception as e:
        return f"Error converting PDF to DOCX: {str(e)}"

# ==================== TEXT EXTRACTION ====================

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX document"""
    try:
        if not DOCX_AVAILABLE:
            return "Error: python-docx library not installed. Run: pip install python-docx"
        
        if not os.path.exists(docx_path):
            return f"File not found: {docx_path}"
        
        doc = Document(docx_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        
        return "\n".join(text_content)
    
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF document"""
    try:
        if not PDF_AVAILABLE:
            return "Error: PDF processing libraries not installed. Run: pip install PyPDF2 pdfplumber"
        
        if not os.path.exists(pdf_path):
            return f"File not found: {pdf_path}"
        
        text_content = []
        
        # Try pdfplumber first (better text extraction)
        try:
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
        except Exception:
            # Fallback to PyPDF2
            try:
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text_content.append(page.extract_text())
            except Exception as e:
                return f"Error extracting text from PDF: {str(e)}"
        
        return "\n".join(text_content)
    
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"

# ==================== TEXT SEARCH ====================

def search_text_in_document(filepath, search_term, case_sensitive=False):
    """Search for text inside a document"""
    try:
        if not os.path.exists(filepath):
            return f"File not found: {filepath}"
        
        file_ext = os.path.splitext(filepath)[1].lower()
        
        if file_ext == '.docx':
            text_content = extract_text_from_docx(filepath)
        elif file_ext == '.pdf':
            text_content = extract_text_from_pdf(filepath)
        else:
            return f"Unsupported file format: {file_ext}"
        
        if text_content.startswith("Error"):
            return text_content
        
        if not case_sensitive:
            text_content = text_content.lower()
            search_term = search_term.lower()
        
        lines = text_content.split('\n')
        matches = []
        
        for line_num, line in enumerate(lines, 1):
            if search_term in line:
                matches.append(f"Line {line_num}: {line.strip()}")
        
        if matches:
            result = f"Found {len(matches)} matches for '{search_term}' in {os.path.basename(filepath)}:\n\n"
            result += "\n".join(matches[:10])  # Show first 10 matches
            if len(matches) > 10:
                result += f"\n... and {len(matches) - 10} more matches"
            return result
        else:
            return f"No matches found for '{search_term}' in {os.path.basename(filepath)}"
    
    except Exception as e:
        return f"Error searching document: {str(e)}"

# ==================== DOCUMENT SUMMARIZATION ====================

def summarize_document(filepath, max_sentences=5):
    """Summarize a document (basic extractive summarization)"""
    try:
        if not os.path.exists(filepath):
            return f"File not found: {filepath}"
        
        file_ext = os.path.splitext(filepath)[1].lower()
        
        if file_ext == '.docx':
            text_content = extract_text_from_docx(filepath)
        elif file_ext == '.pdf':
            text_content = extract_text_from_pdf(filepath)
        else:
            return f"Unsupported file format: {file_ext}"
        
        if text_content.startswith("Error"):
            return text_content
        
        # Basic extractive summarization
        sentences = []
        for paragraph in text_content.split('\n'):
            if paragraph.strip():
                # Split by sentence endings
                para_sentences = paragraph.replace('!', '.').replace('?', '.').split('.')
                for sentence in para_sentences:
                    sentence = sentence.strip()
                    if len(sentence) > 20:  # Filter out very short sentences
                        sentences.append(sentence)
        
        if not sentences:
            return "No content found to summarize"
        
        # Simple scoring based on sentence length and position
        scored_sentences = []
        for i, sentence in enumerate(sentences):
            score = len(sentence.split())  # Word count
            if i < len(sentences) * 0.3:  # Boost early sentences
                score *= 1.5
            scored_sentences.append((score, sentence))
        
        # Sort by score and take top sentences
        scored_sentences.sort(reverse=True)
        top_sentences = [sentence for _, sentence in scored_sentences[:max_sentences]]
        
        summary = f"Summary of {os.path.basename(filepath)} ({max_sentences} key sentences):\n\n"
        summary += "\n\n".join(f"• {sentence}." for sentence in top_sentences)
        
        # Add document stats
        word_count = len(text_content.split())
        summary += f"\n\nDocument Statistics:\n"
        summary += f"• Total words: {word_count}\n"
        summary += f"• Total sentences: {len(sentences)}\n"
        summary += f"• Summary ratio: {max_sentences}/{len(sentences)} sentences"
        
        return summary
    
    except Exception as e:
        return f"Error summarizing document: {str(e)}"

# ==================== DOCUMENT MANAGEMENT ====================

def list_documents(directory=None, file_types=None):
    """List documents in a directory"""
    try:
        if directory is None:
            directory = os.path.join(os.path.expanduser("~"), "Documents")
            if not os.path.exists(directory):
                directory = os.path.expanduser("~")
        
        if file_types is None:
            file_types = ['.docx', '.pdf', '.doc', '.txt', '.rtf']
        
        documents = []
        for file in os.listdir(directory):
            file_path = os.path.join(directory, file)
            if os.path.isfile(file_path):
                file_ext = os.path.splitext(file)[1].lower()
                if file_ext in file_types:
                    file_size = os.path.getsize(file_path)
                    file_size_mb = file_size / (1024 * 1024)
                    modified_time = datetime.fromtimestamp(os.path.getmtime(file_path))
                    documents.append({
                        'name': file,
                        'size_mb': file_size_mb,
                        'modified': modified_time.strftime('%Y-%m-%d %H:%M'),
                        'path': file_path
                    })
        
        if not documents:
            return f"No documents found in {directory}"
        
        # Sort by modification time (newest first)
        documents.sort(key=lambda x: x['modified'], reverse=True)
        
        result = f"Documents in {directory}:\n\n"
        for doc in documents:
            result += f"📄 {doc['name']}\n"
            result += f"   Size: {doc['size_mb']:.2f} MB | Modified: {doc['modified']}\n\n"
        
        return result
    
    except Exception as e:
        return f"Error listing documents: {str(e)}"

def get_document_info(filepath):
    """Get detailed information about a document"""
    try:
        if not os.path.exists(filepath):
            return f"File not found: {filepath}"
        
        file_stats = os.stat(filepath)
        file_size = file_stats.st_size
        created_time = datetime.fromtimestamp(file_stats.st_ctime)
        modified_time = datetime.fromtimestamp(file_stats.st_mtime)
        
        info = f"Document Information: {os.path.basename(filepath)}\n\n"
        info += f"📁 Full Path: {filepath}\n"
        info += f"📏 File Size: {file_size / (1024 * 1024):.2f} MB ({file_size:,} bytes)\n"
        info += f"📅 Created: {created_time.strftime('%Y-%m-%d %H:%M:%S')}\n"
        info += f"📝 Modified: {modified_time.strftime('%Y-%m-%d %H:%M:%S')}\n"
        
        file_ext = os.path.splitext(filepath)[1].lower()
        
        if file_ext in ['.docx', '.pdf']:
            # Try to get text content stats
            if file_ext == '.docx':
                text_content = extract_text_from_docx(filepath)
            else:
                text_content = extract_text_from_pdf(filepath)
            
            if not text_content.startswith("Error"):
                word_count = len(text_content.split())
                char_count = len(text_content)
                line_count = len(text_content.split('\n'))
                
                info += f"\nContent Statistics:\n"
                info += f"📊 Words: {word_count:,}\n"
                info += f"📊 Characters: {char_count:,}\n"
                info += f"📊 Lines: {line_count:,}\n"
        
        return info
    
    except Exception as e:
        return f"Error getting document info: {str(e)}"

# ==================== INSTALLATION CHECK ====================

def check_dependencies():
    """Check which document processing libraries are available"""
    status = "Document Processing Dependencies:\n\n"
    
    if DOCX_AVAILABLE:
        status += "✅ python-docx: Available (Word document support)\n"
    else:
        status += "❌ python-docx: Not installed (pip install python-docx)\n"
    
    if PDF_AVAILABLE:
        status += "✅ PDF libraries: Available (PDF support)\n"
    else:
        status += "❌ PDF libraries: Not installed (pip install PyPDF2 pdfplumber)\n"
    
    if DOCX2PDF_AVAILABLE:
        status += "✅ docx2pdf: Available (DOCX to PDF conversion)\n"
    else:
        status += "❌ docx2pdf: Not installed (pip install docx2pdf)\n"
    
    # Check LibreOffice
    try:
        if platform.system().lower() == "windows":
            subprocess.run(["soffice", "--version"], capture_output=True, check=True)
        else:
            subprocess.run(["libreoffice", "--version"], capture_output=True, check=True)
        status += "✅ LibreOffice: Available (Document conversion support)\n"
    except (subprocess.CalledProcessError, FileNotFoundError):
        status += "❌ LibreOffice: Not installed (Enhanced conversion support)\n"
    
    return status
